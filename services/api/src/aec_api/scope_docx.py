"""SCOPE-DOCX — Exhibit A as an editable Word document.

WHY A SECOND FORMAT AT ALL. The PDF is the signed instrument; this is the one a subcontractor marks
up and sends back. Scope negotiation is redlining, and a PDF is the wrong object for it — the exchange
that actually happens is "here is our scope, strike what you don't hold", which needs a document the
other side can edit. Producing it here rather than letting each GC retype the exhibit is the point:
retyped scope is where clauses quietly go missing.

ONE AUTHORITY FOR WHAT GOES IN. Selection comes from `scope_library.exhibit_clauses`, the same call
the PDF renderer and the `/scope-library/exhibit` preview make. This is deliberately not a third
filter: two copies of that rule is exactly how a plumbing subcontract came to print 31 clauses into a
signed PDF against 11 in its preview, 20 of them duplicated from the agreement body.

NUMBERING IS PRINTED AS LITERAL TEXT, not as Word's own list numbering. An exhibit is argued clause by
clause — "we agreed to 3.2" has to point at exactly one sentence — and Word's automatic numbering
renumbers on edit, which is the one behaviour that must not happen to a cited clause. The label comes
from `scope_library.numbered`, so the DOCX, the PDF, the preview and the JSON all carry identical
labels.

WHY `python-docx` RATHER THAN HAND-ROLLED OOXML. A `.docx` is a zip of XML and this document is only
headings and paragraphs, so writing the parts directly would add no dependency and was seriously
considered. It is refused because this is a **contract document**: a malformed part that Word declines
to open is a much worse failure than a dependency, and it lands on a subcontractor rather than on us.
MIT, approved by the user on 2026-08-07, recorded in `docs/ATTRIBUTIONS.md`.
"""
from __future__ import annotations

import io
from typing import Any

from . import scope_library as sl


def build(ctx: dict[str, Any], clause_ids: list[str] | None = None) -> bytes:
    """Render Exhibit A for `ctx` (project / vendor / trade / merge fields) as .docx bytes.

    `ctx` is the same merge context `contracts.render` builds, so `{{token}}` substitution behaves
    identically in both formats — a clause whose wording differs between the Word copy and the signed
    PDF is a dispute waiting to happen, and the only defence is that both call `sl.merge`.
    """
    from docx import Document  # noqa: PLC0415 — optional-ish, heavy
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Exhibit A — Scope of Work", level=0)

    sub = doc.add_paragraph()
    sub.add_run(f"{ctx.get('project', '')} · {ctx.get('vendor', '')}").italic = True
    if (div := sl.division_for_trade(ctx.get("trade"))):
        doc.add_paragraph(f"CSI Division {div} — {sl.division_name(div)}")

    clauses = sl.numbered(sl.exhibit_clauses(clause_ids, ctx.get("trade")))
    if not clauses:
        # Said out loud rather than producing a document that merely looks short. An empty exhibit
        # attached to a subcontract reads as "no scope was agreed", which is a very different claim
        # from "the composer selected nothing".
        doc.add_paragraph("No scope clauses selected.")
        return _to_bytes(doc)

    # Grouped by category so exclusions get their own visible section. An exclusion buried among
    # inclusions is functionally unstated, and the exclusions are the half this library exists for.
    seen: set[str] = set()
    for c in clauses:
        cat = c["category"]
        if cat not in seen:
            seen.add(cat)
            doc.add_heading(cat, level=1)
        p = doc.add_paragraph()
        p.add_run(f"{c['number']}  {sl.merge(c['title'], ctx)}").bold = True
        body = doc.add_paragraph(sl.merge(c["body"], ctx))
        body.paragraph_format.space_after = Pt(8)

    counts = ", ".join(f"{sum(1 for c in clauses if c['category'] == k)} {k.lower()}"
                       for k in dict.fromkeys(c["category"] for c in clauses))
    tail = doc.add_paragraph()
    tail.add_run(f"This exhibit contains {counts}.").italic = True
    return _to_bytes(doc)


def _to_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
